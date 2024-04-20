import streamlit as st
from Final_file import FlairRecognizer
from Final_file import FlairRecognizer2
import os
import PyPDF2
import docx
# from io import BytesIO
from fpdf import FPDF
import io
from docx import Document
from PiiMaskingService import PiiMaskingService

# Cache the model loading and prediction function
@st.cache_resource
def cached_predict_ner_tags(text):
    return FlairRecognizer.predict_ner_tags(text)

# Cache the text analysis function
@st.cache_resource
def cached_analyze_text(text, operator):
    return FlairRecognizer.analyze_text(text)

@st.cache_resource
def cached_anonimize_text(text, operator):
    return FlairRecognizer2.anonymize(text, operator)

@st.cache_resource
def anonymize(text, operator, model):
    return PiiMaskingService().anonymize(text, operator, model)

def download_masked_file(masked_text, file_extension):
    
    # Create a temporary file to store the masked text
    temp_file_path = f"masked_output.{file_extension}"
    with open(temp_file_path, "w") as temp_file:
        temp_file.write(masked_text)

    # Display a download button
    st.download_button("Download Masked File", temp_file_path, file_name=f"masked_output.{file_extension}")

    # Clean up the temporary file
    os.remove(temp_file_path)

def extract_text_from_pdf(file_contents):
    try:
        # base64_pdf = base64.b64encode(file_contents.read()).decode('utf-8')
        pdf_reader = PyPDF2.PdfReader(file_contents)
        text = ''
        for page_num in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page_num].extract_text()
        return text
    except Exception as e:
        return f"Error occurred: {str(e)}"



def create_pdf(text_content):
    pdf = FPDF()
    pdf.add_page()
    pdf.add_font("DejaVuSans", "", "DejaVuSans.ttf",uni=True)  # Add DejaVuSans font
    pdf.set_font("DejaVuSans", size=12)
    pdf.multi_cell(0, 10, txt=text_content)
    return pdf

def create_word_file(text_content):
    doc = Document()
    doc.add_paragraph(text_content)
    # Save the document to a BytesIO object
    doc_io = io.BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def main():
    st.title('PII Masking App')
    st.sidebar.header('Upload Options')
    upload_option = st.sidebar.radio("Choose upload option:", ('Text Input', 'File Upload'))

    st_operator = st.sidebar.selectbox(
        "De-identification approach",
        ["redact", "replace", "encrypt", "hash", "mask"],
        index=1,
        help="""
        Select which manipulation to the text is requested after PII has been identified.\n
        - Redact: Completely remove the PII text\n
        - Replace: Replace the PII text with a constant, e.g. <PERSON>\n
        - Highlight: Shows the original text with PII highlighted in colors\n
        - Mask: Replaces a requested number of characters with an asterisk (or other mask character)\n
        - Hash: Replaces with the hash of the PII string\n
        - Encrypt: Replaces with an AES encryption of the PII string, allowing the process to be reversed
            """,
    )

    st_model = st.sidebar.selectbox(
        "NER model package",
        [
            "flair/ner-english-large",
            "HuggingFace/obi/deid_roberta_i2b2",
        ],
        index=0,
    )
    
    masked_text_public = ''
    if upload_option == 'Text Input':
        input_text = st.text_area("Enter text here:")
        if st.button('Analyze'):
            with st.spinner('Wait for it... the model is loading'):
                # cached_predict_ner_tags(input_text)
                masked_text = anonymize(input_text, st_operator, st_model)
                # masked_text = cached_anonimize_text(input_text, st_operator)
            st.text_area("Masked text:", value=masked_text, height=200)
    elif upload_option == 'File Upload':
        uploaded_file = st.file_uploader("Upload a file", type=['txt', 'pdf', 'docx'])
        if uploaded_file is not None:
            file_contents = uploaded_file.read()
            #  Process PDF file
            if uploaded_file.type == 'application/pdf':
                extracted_text = extract_text_from_pdf(uploaded_file)
                if st.button('Analyze'):
                    with st.spinner('Wait for it... the model is loading'):
                        # cached_predict_ner_tags(extracted_text)
                        masked_text = anonymize(extracted_text, st_operator, st_model)
                        # masked_text = cached_analyze_text(extracted_text)
                    st.text_area("Masked text:", value=masked_text, height=200) # Display the extracted text
                    if extracted_text:
                        pdf = create_pdf(masked_text)
                        # Save PDF to temporary location
                        pdf_file_path = "masked_output.pdf"
                        pdf.output(pdf_file_path)

                        # Download button
                        st.download_button(label="Download", data=open(pdf_file_path, "rb"), file_name="masked_output.pdf", mime="application/pdf")
                    else:
                        st.warning("Please enter some text to download as PDF.")         
                    
            # Process Word document
            elif uploaded_file.type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                doc = docx.Document(io.BytesIO(file_contents))
                text = ''
                for paragraph in doc.paragraphs:
                    text += paragraph.text
                if st.button('Analyze'):
                    with st.spinner('Wait for it... the model is loading'):
                        # cached_predict_ner_tags(text)
                        masked_text = anonymize(text, st_operator, st_model)
                        # masked_text = cached_analyze_text(text)
                    st.text_area("Masked text:", value=masked_text, height=200)
                    #create word file
                    doc_io = create_word_file(masked_text)
                    #download it
                    st.download_button(label="Download", data=doc_io, file_name="masked_text.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
            else:
                if st.button('Analyze'):
                    with st.spinner('Wait for it... the model is loading'):
                        # cached_predict_ner_tags(file_contents.decode())
                        # masked_text = cached_analyze_text(file_contents.decode())
                        masked_text = anonymize(file_contents.decode(), st_operator, st_model)
                    st.text_area("Masked text:", value=masked_text, height=200)
                    st.download_button(label="Download",data = masked_text,file_name="masked_text.txt")


if __name__ == "__main__":
    main()